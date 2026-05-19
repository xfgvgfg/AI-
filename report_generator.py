import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_markdown_report(ai_markdown: str, original_text: str) -> str:
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    footer = f"""

---

<details>
<summary>附：原始投稿内容（点击展开）</summary>

```
{original_text}
```

</details>

*报告生成时间：{now}*
*由 AI 辅助投稿分析工具自动生成*"""

    return ai_markdown.rstrip() + footer


def export_to_txt(report: str, filename: str = None) -> str:
    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"投稿深度分析报告_{timestamp}.txt"

    filepath = os.path.join(os.path.dirname(os.path.abspath(__file__)), "exports", filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)

    clean = re.sub(r'<details>.*?</details>', '', report, flags=re.DOTALL)
    clean = re.sub(r'<[^>]+>', '', clean)
    clean = re.sub(r'\n{3,}', '\n\n', clean)
    clean = clean.strip()

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(clean)

    return filepath


def export_to_word(report: str, filename: str = None) -> str:
    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"投稿深度分析报告_{timestamp}.docx"

    filepath = os.path.join(os.path.dirname(os.path.abspath(__file__)), "exports", filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)

    title = doc.add_heading('社区投稿深度分析报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    sections = report.split('## ')
    for section in sections[1:]:
        lines = section.strip().split('\n')
        if not lines:
            continue

        heading_text = lines[0].strip().rstrip(':')
        doc.add_heading(heading_text, level=1)

        in_table = False
        table_lines = []

        for line in lines[1:]:
            line = line.strip()
            if not line:
                if in_table and table_lines:
                    _add_table_to_doc(doc, table_lines)
                    table_lines = []
                    in_table = False
                continue

            if line.startswith('|') and line.endswith('|'):
                in_table = True
                table_lines.append(line)
            else:
                if in_table and table_lines:
                    _add_table_to_doc(doc, table_lines)
                    table_lines = []
                    in_table = False

                if line.startswith('> '):
                    p = doc.add_paragraph()
                    run = p.add_run(line.lstrip('> '))
                    run.bold = True
                    run.font.color.rgb = RGBColor(200, 60, 30)
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip('*'))
                    run.bold = True
                    run.font.size = Pt(12)
                elif line.startswith('`'):
                    p = doc.add_paragraph(line.strip('`'))
                    p.style = doc.styles['Normal']
                elif line == '#':
                    continue
                elif line.startswith('# ') and len(line) > 2:
                    doc.add_heading(line[2:], level=2)
                elif '---' in line:
                    doc.add_paragraph('─' * 50)
                elif line.startswith('- '):
                    doc.add_paragraph(line.lstrip('- '), style='List Bullet')
                elif line.startswith('*报告生成') or line.startswith('*由 AI'):
                    p = doc.add_paragraph()
                    run = p.add_run(line.lstrip('* '))
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(128, 128, 128)
                else:
                    doc.add_paragraph(line)

        if in_table and table_lines:
            _add_table_to_doc(doc, table_lines)

    doc.save(filepath)
    return filepath


def _add_table_to_doc(doc, table_lines):
    rows_data = []
    for tl in table_lines:
        tl = tl.strip().strip('|')
        cells = [c.strip() for c in tl.split('|')]
        rows_data.append(cells)

    if not rows_data:
        return

    if len(rows_data) == 1:
        doc.add_paragraph(" | ".join(rows_data[0]))
        return

    is_separator = all('-' in c for c in rows_data[0] if c)
    header_row = None
    data_start = 0
    if is_separator:
        data_start = 1
    else:
        header_row = rows_data[0]
        if len(rows_data) > 1:
            is_sep2 = all('-' in c for c in rows_data[1] if c)
            if is_sep2:
                data_start = 2
            else:
                data_start = 1
        else:
            data_start = 1

    data_rows = rows_data[data_start:]
    if not data_rows:
        return

    num_cols = max(len(r) for r in data_rows)
    if num_cols < 2:
        for r in data_rows:
            doc.add_paragraph(" | ".join(r))
        return

    table = doc.add_table(rows=len(data_rows) + (1 if header_row else 0), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    if header_row:
        for i, cell_text in enumerate(header_row):
            if i < num_cols:
                cell = table.rows[0].cells[i]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    for r_idx, row in enumerate(data_rows):
        doc_row_idx = r_idx + (1 if header_row else 0)
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                table.rows[doc_row_idx].cells[c_idx].text = cell_text

    doc.add_paragraph('')


def get_report_filename(export_format: str = "txt") -> str:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    ext = "docx" if export_format == "word" else "txt"
    return f"投稿深度分析报告_{timestamp}.{ext}"
